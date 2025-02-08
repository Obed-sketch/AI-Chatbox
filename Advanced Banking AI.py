# -*- coding: utf-8 -*-
"""
Banking AI Assistant with Document Export Capabilities
"""

# %% [1] Install Dependencies
!pip install transformers torch flask pandas python-docx reportlab fpdf2 spacy python-docx
!python -m spacy download en_core_web_sm

# %% [2] Import Libraries
import torch
import spacy
from transformers import AutoTokenizer, AutoModelForSequenceClassification, pipeline
from flask import Flask, request, jsonify, render_template, session
from werkzeug.security import generate_password_hash, check_password_hash
import pandas as pd
from docx import Document
from reportlab.pdfgen import canvas
from fpdf import FPDF
import os
import datetime
import hashlib
from functools import wraps

# %% [3] Initialize Core AI Components
class BankingAI:
    def __init__(self):
        # Load financial NLP model
        self.nlp = spacy.load("en_core_web_sm")
        self.tokenizer = AutoTokenizer.from_pretrained("bert-base-uncased")
        self.model = AutoModelForSequenceClassification.from_pretrained("bert-base-uncased")
        self.generator = pipeline("text-generation", model="gpt2")
        
        # Financial knowledge base
        self.knowledge_base = {
            "account_types": ["Checking", "Savings", "CD", "IRA"],
            "loan_products": ["Mortgage", "Personal", "Auto", "Business"],
            "interest_rates": {"Checking": 0.01, "Savings": 0.03, "Mortgage": 0.045}
        }
    
    def analyze_query(self, text):
        """Process user query with multiple NLP layers"""
        doc = self.nlp(text)
        entities = [(ent.text, ent.label_) for ent in doc.ents]
        
        # Classify intent
        inputs = self.tokenizer(text, return_tensors="pt")
        outputs = self.model(**inputs)
        intent = torch.argmax(outputs.logits).item()
        
        # Generate response
        response = self.generator(
            f"Banking query: {text}\nResponse:",
            max_length=100,
            num_return_sequences=1
        )[0]['generated_text']
        
        return {
            "entities": entities,
            "intent": intent,
            "response": response.split("Response:")[-1].strip()
        }

# %% [4] Document Export System
class DocumentExporter:
    @staticmethod
    def to_pdf(data, filename):
        """Export data to PDF"""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        if isinstance(data, dict):
            for key, value in data.items():
                pdf.cell(200, 10, txt=f"{key}: {value}", ln=1)
        elif isinstance(data, pd.DataFrame):
            pdf.cell(200, 10, txt=data.to_string(), ln=1)
        
        pdf.output(filename)
    
    @staticmethod
    def to_word(data, filename):
        """Export data to Word document"""
        doc = Document()
        doc.add_heading('Banking Report', 0)
        
        if isinstance(data, dict):
            for key, value in data.items():
                doc.add_paragraph(f"{key}: {value}")
        elif isinstance(data, pd.DataFrame):
            table = doc.add_table(rows=1, cols=len(data.columns))
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(data.columns):
                hdr_cells[i].text = col
            for _, row in data.iterrows():
                row_cells = table.add_row().cells
                for i, item in enumerate(row):
                    row_cells[i].text = str(item)
        
        doc.save(filename)
    
    @staticmethod
    def to_excel(data, filename):
        """Export data to Excel"""
        if isinstance(data, dict):
            df = pd.DataFrame.from_dict(data, orient='index').reset_index()
            df.columns = ['Category', 'Value']
            df.to_excel(filename, index=False)
        elif isinstance(data, pd.DataFrame):
            data.to_excel(filename, index=False)

# %% [5] Backtesting Framework
class Backtester:
    def __init__(self, ai_model):
        self.ai = ai_model
        self.test_cases = [
            ("What's my account balance?", "balance_inquiry"),
            ("Apply for a mortgage", "loan_application"),
            ("Recent transactions", "transaction_history")
        ]
    
    def run_tests(self):
        results = []
        for query, expected_intent in self.test_cases:
            start_time = time.time()
            response = self.ai.analyze_query(query)
            latency = time.time() - start_time
            
            results.append({
                "query": query,
                "expected_intent": expected_intent,
                "detected_intent": response['intent'],
                "accuracy": 1 if response['intent'] == expected_intent else 0,
                "latency": latency,
                "response_length": len(response['response'])
            })
        return pd.DataFrame(results)

# %% [6] Flask Web Application
app = Flask(__name__)
app.secret_key = os.urandom(24)
ai_system = BankingAI()
exporter = DocumentExporter()

# Security decorator
def requires_auth(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user' not in session:
            return jsonify({"error": "Authentication required"}), 401
        return f(*args, **kwargs)
    return decorated

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/query', methods=['POST'])
@requires_auth
def handle_query():
    data = request.json
    response = ai_system.analyze_query(data['query'])
    
    # Audit log
    audit_data = {
        "user": session['user'],
        "query": data['query'],
        "timestamp": datetime.datetime.now().isoformat()
    }
    pd.DataFrame([audit_data]).to_csv('audit_log.csv', mode='a', header=False)
    
    return jsonify(response)

@app.route('/export', methods=['POST'])
@requires_auth
def export_data():
    data = request.json
    format_type = data['format']
    filename = f"export_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}.{format_type}"
    
    if format_type == 'pdf':
        exporter.to_pdf(data['content'], filename)
    elif format_type == 'docx':
        exporter.to_word(data['content'], filename)
    elif format_type == 'xlsx':
        exporter.to_excel(pd.DataFrame(data['content']), filename)
    
    return jsonify({"status": "success", "filename": filename})

# %% [7] Deployment Setup
if __name__ == "__main__":
    # Initialize components
    backtester = Backtester(ai_system)
    test_results = backtester.run_tests()
    test_results.to_excel("backtest_results.xlsx", index=False)
    
    # Start web server
    app.run(host='0.0.0.0', port=5000, ssl_context='adhoc')
